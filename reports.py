"""
Credit Pulse - Reports (Excel / PDF / Word)  — plain sqlite3
"""
import os, json
from datetime import datetime
from flask import Blueprint, send_file, redirect, url_for
from app import get_db
from utils.translations import get_lang

reports_bp = Blueprint('reports', __name__)
REPORT_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'reports')
os.makedirs(REPORT_DIR, exist_ok=True)

import functools
from flask import session

def login_required(f):
    @functools.wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('admin_logged_in'):
            return redirect(url_for('admin.login'))
        return f(*args, **kwargs)
    return decorated

def _get_data(db, lang):
    survey = db.execute("SELECT * FROM surveys WHERE is_active=1 LIMIT 1").fetchone()
    survey = dict(survey) if survey else {}
    all_resp = [dict(r) for r in db.execute("SELECT * FROM survey_responses").fetchall()]
    departments = {d['id']: dict(d) for d in db.execute("SELECT * FROM departments").fetchall()}

    for r in all_resp:
        rows = db.execute("""SELECT a.answer_value FROM answers a
            JOIN questions q ON q.id=a.question_id
            WHERE a.response_id=? AND q.question_type='rating'""", (r['id'],)).fetchall()
        vals = [float(row['answer_value']) for row in rows if row['answer_value']]
        r['score'] = round(sum(vals)/len(vals), 2) if vals else None

    return survey, all_resp, departments

def _dept_summary(all_resp, departments, lang):
    summary = {}
    for r in all_resp:
        d = departments.get(r['department_id'])
        if not d:
            continue
        name = d['name_ar'] if lang == 'ar' else d['name_en']
        if name not in summary:
            summary[name] = {'count': 0, 'scores': []}
        summary[name]['count'] += 1
        if r['score']:
            summary[name]['scores'].append(r['score'])
    return summary


@reports_bp.route('/excel')
@login_required
def excel_report():
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    from openpyxl.utils import get_column_letter

    db = get_db()
    lang = get_lang()
    survey, all_resp, departments = _get_data(db, lang)
    is_ar = lang == 'ar'

    wb = Workbook()
    NAVY = "1B3A6B"
    GOLD = "C8A951"
    hfont = Font(color="FFFFFF", bold=True, size=11)
    hfill = PatternFill("solid", fgColor=NAVY)
    center = Alignment(horizontal='center')

    # ── Sheet 1: Summary ──
    ws = wb.active
    ws.title = 'ملخص' if is_ar else 'Summary'
    ws.merge_cells('A1:E1')
    ws['A1'] = 'Credit Pulse — تقرير الاستبيان' if is_ar else 'Credit Pulse — Survey Report'
    ws['A1'].font = Font(color="FFFFFF", bold=True, size=15)
    ws['A1'].fill = PatternFill("solid", fgColor=NAVY)
    ws['A1'].alignment = center
    ws.row_dimensions[1].height = 32

    scores = [r['score'] for r in all_resp if r['score']]
    avg = round(sum(scores)/len(scores), 2) if scores else 0
    dept_count = len(set(r['department_id'] for r in all_resp if r['department_id']))
    pulse = round((avg/5)*100, 1)

    summary_rows = [
        ('تاريخ التقرير' if is_ar else 'Report Date', datetime.now().strftime('%Y-%m-%d %H:%M')),
        ('إجمالي المشاركين' if is_ar else 'Total Participants', len(all_resp)),
        ('الأقسام المشاركة' if is_ar else 'Participating Departments', dept_count),
        ('متوسط الرضا' if is_ar else 'Avg Satisfaction', f'{avg} / 5'),
        ('نقاط Credit Pulse' if is_ar else 'Credit Pulse Score', f'{pulse}%'),
    ]
    for r_num, (label, val) in enumerate(summary_rows, 3):
        ws.cell(r_num, 1, label).font = Font(bold=True)
        ws.cell(r_num, 2, str(val))
    ws.column_dimensions['A'].width = 32
    ws.column_dimensions['B'].width = 22

    # ── Sheet 2: Responses ──
    ws2 = wb.create_sheet('الإجابات' if is_ar else 'Responses')
    heads = ['#', 'القسم' if is_ar else 'Department',
             'الاسم' if is_ar else 'Name',
             'الرقم الوظيفي' if is_ar else 'Employee ID',
             'مجهول' if is_ar else 'Anonymous',
             'متوسط التقييم' if is_ar else 'Avg Rating',
             'تاريخ الإرسال' if is_ar else 'Submitted']
    for ci, h in enumerate(heads, 1):
        c = ws2.cell(1, ci, h)
        c.font = hfont; c.fill = hfill; c.alignment = center
    for ri, r in enumerate(all_resp, 2):
        d = departments.get(r['department_id'])
        dn = (d['name_ar'] if is_ar else d['name_en']) if d else '—'
        ws2.append([ri-1, dn,
                    r['employee_name'] or ('—' if r['is_anonymous'] else ''),
                    r['employee_id'] or ('—' if r['is_anonymous'] else ''),
                    'نعم' if r['is_anonymous'] else 'لا' if is_ar else ('Yes' if r['is_anonymous'] else 'No'),
                    r['score'] or '—',
                    str(r['submitted_at'])[:16] if r['submitted_at'] else ''])
    for ci in range(1, len(heads)+1):
        ws2.column_dimensions[get_column_letter(ci)].width = 22

    # ── Sheet 3: Dept Analysis ──
    ws3 = wb.create_sheet('تحليل الأقسام' if is_ar else 'Dept Analysis')
    dheads = ['القسم' if is_ar else 'Department',
              'المشاركون' if is_ar else 'Participants',
              'متوسط الرضا' if is_ar else 'Avg Rating',
              'Pulse Score']
    for ci, h in enumerate(dheads, 1):
        c = ws3.cell(1, ci, h)
        c.font = hfont; c.fill = hfill; c.alignment = center
    for name, info in _dept_summary(all_resp, departments, lang).items():
        d_avg = round(sum(info['scores'])/len(info['scores']), 2) if info['scores'] else 0
        ws3.append([name, info['count'], d_avg, f"{round((d_avg/5)*100,1)}%"])
    for ci in range(1, len(dheads)+1):
        ws3.column_dimensions[get_column_letter(ci)].width = 28

    path = os.path.join(REPORT_DIR, f'CreditPulse_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx')
    wb.save(path)
    return send_file(path, as_attachment=True, download_name='CreditPulse_Report.xlsx')


@reports_bp.route('/pdf')
@login_required
def pdf_report():
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT

    db = get_db()
    lang = get_lang()
    is_ar = lang == 'ar'
    survey, all_resp, departments = _get_data(db, lang)

    path = os.path.join(REPORT_DIR, f'CreditPulse_{datetime.now().strftime("%Y%m%d_%H%M%S")}.pdf')
    doc = SimpleDocTemplate(path, pagesize=A4,
        rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)

    styles = getSampleStyleSheet()
    NAVY = colors.HexColor('#1B3A6B')
    GOLD = colors.HexColor('#C8A951')
    align = TA_RIGHT if is_ar else TA_LEFT

    T = lambda text, **kw: Paragraph(text, ParagraphStyle('x', parent=styles['Normal'], **kw))

    story = []
    story.append(Spacer(1, 0.5*cm))
    story.append(T("CREDIT PULSE", fontSize=24, fontName='Helvetica-Bold',
                   textColor=NAVY, alignment=TA_CENTER, spaceAfter=6))
    story.append(T(
        "تقرير استبيان قطاع الائتمان" if is_ar else "Credit Sector Survey Report",
        fontSize=14, fontName='Helvetica-Bold', textColor=GOLD, alignment=TA_CENTER, spaceAfter=4))
    story.append(T("البنك الأهلي المصري | National Bank of Egypt",
                   fontSize=11, alignment=TA_CENTER, spaceAfter=6, textColor=colors.grey))
    story.append(HRFlowable(width="100%", thickness=2, color=GOLD, spaceAfter=12))

    scores = [r['score'] for r in all_resp if r['score']]
    avg = round(sum(scores)/len(scores), 2) if scores else 0
    pulse = round((avg/5)*100, 1)
    dept_count = len(set(r['department_id'] for r in all_resp if r['department_id']))

    story.append(T("الملخص التنفيذي" if is_ar else "Executive Summary",
                   fontSize=14, fontName='Helvetica-Bold', textColor=NAVY, spaceAfter=8))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=8))

    summ = [
        ['إجمالي المشاركين' if is_ar else 'Total Participants', str(len(all_resp))],
        ['الأقسام المشاركة' if is_ar else 'Participating Departments', str(dept_count)],
        ['متوسط الرضا' if is_ar else 'Avg. Satisfaction', f'{avg} / 5.0'],
        ['نقاط Credit Pulse' if is_ar else 'Credit Pulse Score', f'{pulse}%'],
    ]
    tbl = Table(summ, colWidths=[10*cm, 6*cm])
    tbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0),(0,-1), colors.HexColor('#EEF2F7')),
        ('FONTNAME', (0,0),(0,-1), 'Helvetica-Bold'),
        ('ALIGN', (1,0),(1,-1), 'CENTER'),
        ('GRID', (0,0),(-1,-1), 0.5, colors.lightgrey),
        ('ROWBACKGROUNDS', (0,0),(-1,-1), [colors.HexColor('#F8F9FA'), colors.white]),
        ('PADDING', (0,0),(-1,-1), 8),
        ('FONTSIZE', (0,0),(-1,-1), 11),
    ]))
    story.append(tbl)
    story.append(Spacer(1, 0.6*cm))

    # Dept table
    story.append(T("تحليل الأقسام" if is_ar else "Department Analysis",
                   fontSize=14, fontName='Helvetica-Bold', textColor=NAVY, spaceAfter=8))
    story.append(HRFlowable(width="100%", thickness=1, color=NAVY, spaceAfter=8))
    dept_rows = [['القسم' if is_ar else 'Department',
                  'المشاركون' if is_ar else 'Participants',
                  'متوسط الرضا' if is_ar else 'Avg Rating',
                  'Pulse Score']]
    for name, info in list(_dept_summary(all_resp, departments, lang).items())[:12]:
        d_avg = round(sum(info['scores'])/len(info['scores']), 2) if info['scores'] else 0
        dept_rows.append([name, str(info['count']), f'{d_avg}/5', f"{round((d_avg/5)*100,1)}%"])
    dtbl = Table(dept_rows, colWidths=[8*cm, 3*cm, 3.5*cm, 3*cm])
    dtbl.setStyle(TableStyle([
        ('BACKGROUND', (0,0),(-1,0), NAVY),
        ('TEXTCOLOR', (0,0),(-1,0), colors.white),
        ('FONTNAME', (0,0),(-1,0), 'Helvetica-Bold'),
        ('ALIGN', (0,0),(-1,-1), 'CENTER'),
        ('GRID', (0,0),(-1,-1), 0.5, colors.lightgrey),
        ('ROWBACKGROUNDS', (0,1),(-1,-1), [colors.HexColor('#F8F9FA'), colors.white]),
        ('PADDING', (0,0),(-1,-1), 7),
        ('FONTSIZE', (0,0),(-1,-1), 10),
    ]))
    story.append(dtbl)
    story.append(Spacer(1, 1*cm))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=6))
    story.append(T(
        "تقرير سري — للاستخدام الداخلي فقط | Credit Pulse © البنك الأهلي المصري" if is_ar
        else "Confidential — Internal Use Only | Credit Pulse © National Bank of Egypt",
        fontSize=9, textColor=colors.grey, alignment=TA_CENTER))

    doc.build(story)
    return send_file(path, as_attachment=True, download_name='CreditPulse_Report.pdf')


@reports_bp.route('/word')
@login_required
def word_report():
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    db = get_db()
    lang = get_lang()
    is_ar = lang == 'ar'
    survey, all_resp, departments = _get_data(db, lang)

    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2); s.bottom_margin = Cm(2)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

    t0 = doc.add_heading('Credit Pulse', 0)
    t0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t0.runs[0].font.color.rgb = RGBColor(0x1B,0x3A,0x6B)

    sub = doc.add_paragraph('تقرير استبيان قطاع الائتمان' if is_ar else 'Credit Sector Survey Report')
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.bold = True
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].font.color.rgb = RGBColor(0xC8,0xA9,0x51)

    bank = doc.add_paragraph('البنك الأهلي المصري | National Bank of Egypt')
    bank.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bank.runs[0].font.size = Pt(11)

    doc.add_paragraph(f"{'تاريخ التقرير' if is_ar else 'Report Date'}: {datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()

    scores = [r['score'] for r in all_resp if r['score']]
    avg = round(sum(scores)/len(scores), 2) if scores else 0
    pulse = round((avg/5)*100, 1)
    dept_count = len(set(r['department_id'] for r in all_resp if r['department_id']))

    h1 = doc.add_heading('الملخص التنفيذي' if is_ar else 'Executive Summary', 1)
    h1.runs[0].font.color.rgb = RGBColor(0x1B,0x3A,0x6B)

    tbl = doc.add_table(4, 2); tbl.style = 'Table Grid'
    for i, (label, val) in enumerate([
        ('إجمالي المشاركين' if is_ar else 'Total Participants', str(len(all_resp))),
        ('الأقسام المشاركة' if is_ar else 'Departments', str(dept_count)),
        ('متوسط الرضا' if is_ar else 'Avg. Satisfaction', f'{avg} / 5.0'),
        ('نقاط Credit Pulse' if is_ar else 'Credit Pulse Score', f'{pulse}%'),
    ]):
        tbl.rows[i].cells[0].text = label
        tbl.rows[i].cells[1].text = val
        tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()
    h2 = doc.add_heading('تحليل الأقسام' if is_ar else 'Department Analysis', 1)
    h2.runs[0].font.color.rgb = RGBColor(0x1B,0x3A,0x6B)

    ds = _dept_summary(all_resp, departments, lang)
    if ds:
        dtbl = doc.add_table(1+len(ds), 4); dtbl.style = 'Table Grid'
        for ci, h in enumerate(['القسم' if is_ar else 'Department',
                                  'المشاركون' if is_ar else 'Participants',
                                  'متوسط الرضا' if is_ar else 'Avg Rating', 'Pulse Score']):
            dtbl.rows[0].cells[ci].text = h
            dtbl.rows[0].cells[ci].paragraphs[0].runs[0].font.bold = True
        for ri, (name, info) in enumerate(ds.items(), 1):
            d_avg = round(sum(info['scores'])/len(info['scores']), 2) if info['scores'] else 0
            for ci, val in enumerate([name, str(info['count']), f'{d_avg}/5', f"{round((d_avg/5)*100,1)}%"]):
                dtbl.rows[ri].cells[ci].text = val

    doc.add_paragraph()
    footer = doc.add_paragraph('تقرير سري للاستخدام الداخلي فقط' if is_ar else 'Confidential — Internal Use Only')
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.color.rgb = RGBColor(0x80,0x80,0x80)

    path = os.path.join(REPORT_DIR, f'CreditPulse_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
    doc.save(path)
    return send_file(path, as_attachment=True, download_name='CreditPulse_Report.docx')
